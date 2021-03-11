"""python PyCertGen.py"""
from docx import Document
import os


def parser(document, v=1):
    """
    v : Verbosity
    """
    LINES = []
    for p in document.paragraphs:
        inline = p.runs
        for i in range(len(inline)):
            text = inline[i].text
            if v == 1:
                print(text)
            LINES.append(text)
    return LINES


def replacer(document, dic):
    for p in document.paragraphs:
        inline = p.runs
        for i in range(len(inline)):
            text = inline[i].text
            if text in dic.keys():
                text = text.replace(text, dic[text])
                inline[i].text = text
    return 1

# For various people load multiple dics


def DocxLoader(FileName):
    return Document(FileName)


if __name__ == '__main__':
    ROOT_DIR = os.getcwd()
    FileName = os.path.join(ROOT_DIR, "CertTemplateSamples",
                            "Certificate_of_Appreciation.docx")
    document = Document(FileName)

    dic = {'YOUR_NAME': 'Farhan Hai Khan',
           'Outstanding_Professional_Experience': 'Machine Learning Engineer',
           'Date_Time': '12th of March 2021',
           'Sign': 'M. Agarwal',
           'Signatory_Name': 'Manish Agarwal',
           'Signatory_Position': 'IT Team Head',
           }
    Text_List = parser(document)
    IsReplaced = replacer(document, dic)
    OutFileName = os.path.join(ROOT_DIR, "TempSaved", "Cert_Farhan.docx")
    document.save(OutFileName)
